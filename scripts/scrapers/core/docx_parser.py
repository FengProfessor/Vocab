import sys
import os
import json
import docx

def parse_docx(file_path):
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}
    
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        return {"error": f"Failed to open document: {str(e)}"}
    
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
            
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Deduplicate paragraphs inside cell
                cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()]).strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)
        
    return {
        "paragraphs": paragraphs,
        "tables": tables
    }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Missing input file path"}))
        sys.exit(1)
        
    file_path = sys.argv[1]
    result = parse_docx(file_path)
    print(json.dumps(result, ensure_ascii=True))
