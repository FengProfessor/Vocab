# -*- coding: utf-8 -*-
"""Export grammar wordbanks quality preview to DOCX."""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("tmp/wordbanks-export.json")
OUT = Path("tmp/Grammar_Wordbanks_Full_All_Levels.docx")

TITLES = {
    "countable-uncountable": "1. Danh từ đếm được & không đếm được (C/U)",
    "plural-nouns": "2. Danh từ số nhiều (Plural)",
    "past-simple": "3. Past Simple · Irregular V1–V2–V3",
    "articles": "4. Mạo từ a / an / the / zero",
    "quantifiers": "5. Lượng từ (some/any/much/many…)",
    "present-simple": "6. Present Simple · ngôi 3 + tần suất",
    "verb-to-be": "7. To be (am/is/are)",
    "possessives": "8. Sở hữu (my/mine · 's)",
    "demonstratives": "9. This / That / These / Those",
    "there-is-there-are": "10. There is / There are",
    "prepositions-place": "11. Giới từ nơi chốn",
    "prepositions-time": "12. Giới từ thời gian",
    "comparatives-superlatives": "13. So sánh hơn / nhất",
    "adverbs-frequency": "14. Trạng từ tần suất",
}

# auto-extend titles from export json
import json as _json
_data=_json.loads(Path('tmp/wordbanks-export.json').read_text(encoding='utf-8'))
for _i,_slug in enumerate(_data.get('stats',{}),1):
    if _slug not in TITLES:
        TITLES[_slug]=f'{_i}. {_slug}'
ORDER = list(TITLES.keys())
# prefer stats order
ORDER = list(_data.get('stats',{}).keys()) or ORDER


def set_cell_shading(cell, hex_color: str):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, rows):
    if not rows:
        return
    keys = list(rows[0].keys())
    table = doc.add_table(rows=1 + len(rows), cols=len(keys))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, k in enumerate(keys):
        hdr[i].text = str(k)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(hdr[i], "1E3A5F")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, k in enumerate(keys):
            cells[ci].text = str(row.get(k, "") or "")
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
        if ri % 2 == 1:
            for ci in range(len(keys)):
                set_cell_shading(cells[ci], "F1F5F9")
    doc.add_paragraph()


def main():
    data = json.loads(SRC.read_text(encoding="utf-8"))
    stats = data["stats"]
    lessons = data["lessons"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    t = doc.add_heading("LingoPro · Grammar Wordbanks — Quality Preview", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Bảng case đặc biệt (format chuyên đề GV VN) · {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        "Nguồn: scripts/grammar-a0a2/wordbanks-dense.mjs → sections.wordbanks (DB live)"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_heading("0. Tổng quan mật độ", level=1)
    doc.add_paragraph(
        "Mỗi bài có 1–7 bảng; mỗi dòng ≈ 1 từ/cặp (không nhồi nhiều từ vào 1 ô). "
        "Mục tiêu: độ chi tiết handout ôn thi (C/U, irregular plurals, V1–V2–V3…)."
    )

    overview = []
    total_tables = 0
    total_rows = 0
    for slug in ORDER:
        st = stats.get(slug, {})
        tables = st.get("tables", 0)
        rows = st.get("rows", 0)
        total_tables += tables
        total_rows += rows
        overview.append(
            {
                "Bài (slug)": slug,
                "Số bảng": str(tables),
                "Số dòng": str(rows),
            }
        )
    overview.append(
        {"Bài (slug)": "TỔNG", "Số bảng": str(total_tables), "Số dòng": str(total_rows)}
    )
    add_table(doc, overview)

    doc.add_heading("Ghi chú review chất lượng", level=2)
    for line in [
        "✓ Cột 'Sai hay gặp / Cách đúng' bám lỗi VN (informations, homeworks, childs…).",
        "✓ Irregular verbs đủ ~130 dạng handout THCS; Past Simple dùng V2.",
        "✓ Plural: regular theo rule + irregular 1 dòng/1 từ + plural-only (scissors…).",
        "✓ C/U: food, abstract, material, sport/health, unitiser, dual-face.",
        "⚠ Bài ngắn (demonstratives, there-is) chủ yếu rule table — chưa list 'trăm từ'.",
        "⚠ Chưa deploy Vercel (nếu prod UI cũ) — data đã trong Supabase.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    for slug in ORDER:
        banks = lessons.get(slug) or []
        if not banks:
            continue
        doc.add_page_break()
        doc.add_heading(TITLES.get(slug, slug), level=1)
        st = stats.get(slug, {})
        meta = doc.add_paragraph()
        meta.add_run(
            f"Slug: {slug} · {st.get('tables', 0)} bảng · {st.get('rows', 0)} dòng"
        ).italic = True

        for bi, bank in enumerate(banks, 1):
            title = bank.get("title") or f"Bảng {bi}"
            icon = bank.get("icon") or ""
            doc.add_heading(f"{icon} {title}".strip(), level=2)
            note = bank.get("note")
            if note:
                np = doc.add_paragraph()
                nr = np.add_run(f"Ghi chú: {note}")
                nr.font.size = Pt(9)
                nr.font.color.rgb = RGBColor(146, 64, 14)
            rows = bank.get("rows") or []
            doc.add_paragraph(f"{len(rows)} dòng").runs[0].font.size = Pt(9)
            add_table(doc, rows)

    doc.add_page_break()
    doc.add_heading("Phụ lục · Cách render trên app", level=1)
    doc.add_paragraph(
        "UI: components/grammar/GoldenLesson.tsx — card tag “Bảng từ”, table sticky header, max-height scroll."
    )
    doc.add_paragraph(
        "Apply: node scripts/grammar-a0a2/apply-a0a2-quality.mjs — gắn banksForSlug(slug) vào sections.wordbanks."
    )
    doc.add_paragraph(
        "File nguồn: scripts/grammar-a0a2/wordbanks-dense.mjs"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"WROTE {OUT.resolve()}")
    print(f"tables={total_tables} rows={total_rows}")


if __name__ == "__main__":
    main()
